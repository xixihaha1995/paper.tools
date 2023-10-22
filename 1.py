'''
filePath = "C:\\Users\wulic\\OneDrive - University of Wyoming\\PhD\\e_Dissertation"
fileName = r"1.摘录摘要引用结论.docx"

Find paragraph with style "Heading 2"
And add the following three lines:
【摘要】
【引入】
【结论】
'''
import re, os
from docx import Document
from docx.oxml.xmlchemy import OxmlElement
from docx.text.paragraph import Paragraph

filePath = r"C:\Users\wulic\OneDrive - University of Wyoming\PhD\e_Dissertation"
fileName = r"1.摘录摘要引用结论.docx"
targetHeading = 'Heading 2'
newContent = ['【摘要】', '【引入】', '【结论】']
newAddedStyle = 'Normal'
# 打开.docx文件
doc = Document(os.path.join(filePath, fileName))
# 遍历所有段落
def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para
for paragraph in doc.paragraphs:
    # 读取段落的样式
    paragraph_style = paragraph.style
    # print('段落样式:', paragraph_style.name)
    if paragraph_style.name == targetHeading:
        # Reversely insert new content after the targetHeading
        for i in range(len(newContent)-1, -1, -1):
            insert_paragraph_after(paragraph, newContent[i], newAddedStyle)



doc.save(os.path.join(filePath, "New_"+fileName))
